from __future__ import annotations

import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt


def insert_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_p.append(deepcopy(paragraph._p.pPr))
    new_para._p = new_p
    run = new_para.add_run(text)
    return new_para, run


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("Usage: add_appendix_e_to_toc.py <docx-path>")

    path = Path(sys.argv[1])
    document = Document(path)

    toc_start = next(
        (idx for idx, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "TABLE OF CONTENTS"),
        None,
    )
    if toc_start is None:
        raise SystemExit("TABLE OF CONTENTS not found")

    toc_end = next(
        (
            idx
            for idx, paragraph in enumerate(document.paragraphs[toc_start + 1 :], start=toc_start + 1)
            if paragraph.text.strip().upper() == "ABSTRACT"
        ),
        len(document.paragraphs),
    )
    toc_paragraphs = document.paragraphs[toc_start:toc_end]
    if any("Appendix E: Important Code Snippets" in paragraph.text for paragraph in toc_paragraphs):
        print("Appendix E already exists in TOC; no changes made.")
        return

    reference_para = next((p for p in toc_paragraphs if p.text.startswith("References\t")), None)
    if reference_para is None:
        raise SystemExit("References entry not found in TOC")

    new_para, run = insert_after(reference_para, "Appendix E: Important Code Snippets\t41")
    new_para.paragraph_format.left_indent = Inches(0)
    new_para.paragraph_format.line_spacing = 1.5
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(17)

    document.save(path)
    print(path)


if __name__ == "__main__":
    main()
