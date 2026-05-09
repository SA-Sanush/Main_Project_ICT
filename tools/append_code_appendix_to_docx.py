from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


CODE_SNIPPETS = [
    (
        "Database Initialization",
        """def init_db() -> None:
    conn = sqlite3.connect(DB_PATH)
    try:
        conn.execute(
            "CREATE TABLE IF NOT EXISTS reports ("
            "id INTEGER PRIMARY KEY AUTOINCREMENT,"
            "user_id INTEGER, filename TEXT, uploaded_at TEXT,"
            "score INTEGER, matched_jobs TEXT, criteria TEXT,"
            "suggestions TEXT, analysis TEXT)"
        )
        conn.commit()
    finally:
        conn.close()""",
    ),
    (
        "Allowed Resume File Validation",
        """def allowed_file(filename: str) -> bool:
    return (
        "." in filename
        and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS
    )""",
    ),
    (
        "Resume Text Extraction",
        """def extract_resume_text(path: Path, extension: str) -> str:
    if extension == "pdf":
        return extract_text_from_pdf(path)
    if extension == "docx":
        return extract_text_from_docx(path)
    if extension == "txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    return """"",
    ),
    (
        "Criteria-wise Score Construction",
        """def build_criteria_scores(text: str, extension: str, profile=None, target_text: str = "") -> dict:
    return {
        "Contact Information": detect_contact_info(text),
        "Professional Summary": detect_summary(text),
        "Work Experience": detect_experience(text, profile),
        "Education": detect_education(text),
        "Skills": detect_skills(text, profile),
        "TAS Optimization": detect_tas_optimization(text),
        "Consistency": detect_consistency(text),
        "Proofreading": detect_proofreading(text),
        "File Format": 100 if extension in ALLOWED_EXTENSIONS else 20,
        "Relevance": detect_relevance(text, target_text),
    }""",
    ),
    (
        "Role-adjusted Final Score",
        """def calculate_role_adjusted_score(criteria_scores: dict, target_title: str = "") -> int:
    weights = scoring_weights_for_role(target_title)
    total = sum(criteria_scores.get(name, 0) * weight for name, weight in weights.items())
    return int(np.clip(round(total), 0, 100))""",
    ),
    (
        "Job Recommendation Match",
        """def find_applicable_jobs(text: str, profile: dict | None = None, target_job: dict | None = None) -> list:
    jobs_df = load_job_data()
    resume_terms = build_weighted_terms(text)
    recommendations = []
    if target_job:
        recommendations.append(target_job)
    for _, row in jobs_df.iterrows():
        keywords = [item.strip() for item in row["keywords"].split(",")]
        job_text = f"{row['title']} {row.get('description', '')} {' '.join(keywords)}"
        similarity = weighted_cosine_similarity(resume_terms, build_weighted_terms(job_text))
        if similarity > 0.18:
            recommendations.append({"title": row["title"], "match_score": round(similarity * 100)})
    return recommendations""",
    ),
    (
        "Admin Catalog Update",
        """new_row = pd.DataFrame([{
    "title": title,
    "keywords": keywords,
    "description": description,
}])
jobs_df = pd.concat([jobs_df, new_row], ignore_index=True)
jobs_df.to_csv(JOB_DATA_PATH, index=False)""",
    ),
]


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_heading("", level=1)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.bold = True
    run.font.size = Pt(24)


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_code_block(document: Document, heading: str, code: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(heading)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    code_paragraph = document.add_paragraph()
    code_paragraph.paragraph_format.left_indent = Inches(0.25)
    code_paragraph.paragraph_format.line_spacing = 1.0
    run = code_paragraph.add_run(code)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(8.5)


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("Usage: append_code_appendix_to_docx.py <docx-path>")
    path = Path(sys.argv[1])
    document = Document(path)
    if any(p.text.strip().upper() == "APPENDIX E: IMPORTANT CODE SNIPPETS" for p in document.paragraphs):
        print("Appendix already exists; no changes made.")
        return

    document.add_page_break()
    add_heading(document, "APPENDIX E: IMPORTANT CODE SNIPPETS")
    add_paragraph(
        document,
        "This appendix includes selected code snippets from the application to show the core implementation approach. "
        "The complete source code is maintained separately in the project folder, while the report includes only the "
        "most relevant parts for academic reference.",
    )
    add_paragraph(
        document,
        "The snippets cover database initialization, resume file validation, text extraction, scoring, role matching "
        "and catalog update logic. These areas represent the main technical workflow of the Talent Acquisition System.",
    )
    for heading, code in CODE_SNIPPETS:
        add_code_block(document, heading, code)
    document.save(path)
    print(path)


if __name__ == "__main__":
    main()
