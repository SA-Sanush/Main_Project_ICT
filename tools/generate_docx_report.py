from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "reports" / "Talent_Acquisition_System_Project_Report.txt"
OUTPUT = ROOT / "reports" / "Talent_Acquisition_System_Project_Report.docx"
CHART = ROOT / "static" / "images" / "score_chart.png"


FIGURES = [
    ("fig 4.1.1", "System architecture", "17"),
    ("fig 11.2.1", "Home page and resume upload form", "Appendix"),
    ("fig 11.2.2", "Login page", "Appendix"),
    ("fig 11.2.3", "Resume analysis result page", "Appendix"),
    ("fig 11.2.4", "Scan history dashboard", "Appendix"),
    ("fig 11.2.5", "Admin dashboard", "Appendix"),
    ("fig 11.2.6", "Contact page", "Appendix"),
]


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_cover(document: Document) -> None:
    for _ in range(4):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TALENT ACQUISITION SYSTEM")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(18)

    lines = [
        "",
        "A project report submitted to ICT Academy of Kerala",
        "in partial fulfillment of the requirements",
        "for the certification of",
        "ADVANCED PYTHON",
        "",
        "submitted by",
        "SASAN",
        "",
        "",
        "ICT ACADEMY OF KERALA",
        "THIRUVANANTHAPURAM, KERALA, INDIA",
        "MAY 2026",
    ]
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        if line in {"ADVANCED PYTHON", "SASAN", "ICT ACADEMY OF KERALA"}:
            run.bold = True
    document.add_page_break()


def add_list_of_figures(document: Document) -> None:
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("LIST OF FIGURES")
    run.bold = True
    run.font.size = Pt(14)

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["SL NO.", "FIGURES", "PAGE NO."]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
    for figure_no, title, page in FIGURES:
        cells = table.add_row().cells
        set_cell_text(cells[0], figure_no)
        set_cell_text(cells[1], title.upper() if title == "System architecture" else title)
        set_cell_text(cells[2], page)
    document.add_page_break()


def add_contents(document: Document) -> None:
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("TABLE OF CONTENTS")
    run.bold = True
    run.font.size = Pt(14)
    rows = [
        ("ABSTRACT", "4"),
        ("1. PROBLEM DEFINITION", "5"),
        ("1.1 OVERVIEW", "5"),
        ("1.2 PROBLEM STATEMENT", "5"),
        ("2. INTRODUCTION", "8"),
        ("3. SYSTEM ANALYSIS", "8"),
        ("3.1 EXISTING SYSTEM", "8"),
        ("3.2 PROPOSED SYSTEM", "8"),
        ("3.3 FEASIBILITY STUDY", "9"),
        ("3.4 TECHNOLOGIES USED", "11"),
        ("3.5 LANGUAGE SPECIFICATIONS", "12"),
        ("4. SYSTEM DESIGN", "16"),
        ("4.1 SYSTEM ARCHITECTURE", "17"),
        ("5. PROJECT DESCRIPTION", "18"),
        ("6. SYSTEM TESTING AND IMPLEMENTATION", "20"),
        ("7. SYSTEM MAINTENANCE", "22"),
        ("8. FUTURE ENHANCEMENTS", "24"),
        ("9. CONCLUSION", "25"),
        ("10. BIBLIOGRAPHY", "26"),
        ("11. APPENDIX", "28"),
        ("11.1 SYSTEM CODING", "28"),
        ("11.2 SCREENSHOTS", "37"),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "CONTENT", True)
    set_cell_text(table.rows[0].cells[1], "PAGE No.", True)
    for label, page in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label)
        set_cell_text(cells[1], page)
    document.add_page_break()


def classify_heading(line: str) -> str | None:
    upper = line.upper()
    main = {
        "ABSTRACT",
        "1. PROBLEM DEFINITION",
        "2. INTRODUCTION",
        "3. SYSTEM ANALYSIS",
        "4. SYSTEM DESIGN",
        "5. PROJECT DESCRIPTION",
        "6. SYSTEM TESTING AND IMPLEMENTATION",
        "7. SYSTEM MAINTENANCE",
        "8. FUTURE ENHANCEMENTS",
        "9. CONCLUSION",
        "10. BIBLIOGRAPHY",
        "11. APPENDIX",
    }
    if upper in main:
        return "Heading 1"
    if line[:4].count(".") >= 1 and any(line.startswith(prefix) for prefix in [f"{i}." for i in range(1, 12)]):
        return "Heading 2"
    return None


def add_body_from_source(document: Document) -> None:
    text = SOURCE.read_text(encoding="utf-8")
    body = text.split("ABSTRACT", 1)[1]
    lines = ["ABSTRACT"] + body.splitlines()
    skip_prefixes = {"LIST OF FIGURES", "TABLE OF CONTENTS"}
    for raw_line in lines:
        line = raw_line.strip()
        if not line or line in skip_prefixes:
            continue
        if line.startswith("SL NO.") or line.startswith("Note: Screenshot placeholders"):
            continue
        heading_style = classify_heading(line)
        if heading_style:
            if heading_style == "Heading 1" and line != "ABSTRACT":
                document.add_page_break()
            paragraph = document.add_paragraph(line.upper(), style=heading_style)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(line[2:], style="List Bullet")
        elif len(line) > 2 and line[0].isdigit() and line[1] == "." and " " in line[:5]:
            paragraph = document.add_paragraph(line, style="List Number")
        else:
            paragraph = document.add_paragraph(line)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def add_appendix_assets(document: Document) -> None:
    document.add_page_break()
    document.add_heading("11.2 SCREENSHOTS", level=2)
    document.add_paragraph(
        "The application screens are generated from the Flask routes. The following figures "
        "correspond to the working pages of the Talent Acquisition System."
    )
    if CHART.exists():
        document.add_paragraph("Fig 11.2.3 Resume analysis score chart")
        document.add_picture(str(CHART), width=Inches(5.5))
    for _, title, _ in FIGURES[1:]:
        paragraph = document.add_paragraph()
        paragraph.add_run(title).bold = True
        document.add_paragraph(
            "Screenshot placeholder: open the running Flask application and capture this page "
            "for final submission if required by the institution."
        )


def main() -> None:
    document = Document()
    configure_document(document)
    add_cover(document)
    add_list_of_figures(document)
    add_contents(document)
    add_body_from_source(document)
    add_appendix_assets(document)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
