from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("Usage: order_appendix_e_before_references.py <docx-path>")
    path = Path(sys.argv[1])
    document = Document(path)

    toc_start = next(
        (idx for idx, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "TABLE OF CONTENTS"),
        None,
    )
    if toc_start is None:
        raise SystemExit("TABLE OF CONTENTS not found")
    toc_end = next(
        (
            idx
            for idx, paragraph in enumerate(document.paragraphs[toc_start + 1 :], start=toc_start + 1)
            if paragraph.text.strip().upper() == "ABSTRACT" and paragraph.style.name.startswith("Heading")
        ),
        len(document.paragraphs),
    )

    entries = []
    for paragraph in document.paragraphs[toc_start + 1 : toc_end]:
        text = paragraph.text.strip()
        if text:
            entries.append(text)

    if "References\t47" in entries and "Appendix E: Important Code Snippets\t45" in entries:
        entries.remove("References\t47")
        appendix_index = entries.index("Appendix E: Important Code Snippets\t45")
        entries.insert(appendix_index + 1, "References\t47")

    paragraphs = document.paragraphs[toc_start + 1 : toc_end]
    for idx, paragraph in enumerate(paragraphs):
        paragraph.clear()
        if idx >= len(entries):
            continue
        text = entries[idx]
        title = text.split("\t", 1)[0]
        first_token = title.split(" ", 1)[0]
        subheading = "." in first_token and first_token.split(".")[-1].isdigit()
        paragraph.paragraph_format.left_indent = Inches(0.28 if subheading else 0)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12 if subheading else 13)
        run.bold = not subheading

    document.save(path)
    print(path)


if __name__ == "__main__":
    main()
