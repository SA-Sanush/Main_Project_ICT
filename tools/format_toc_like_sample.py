from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def is_subheading(title: str) -> bool:
    first = title.split(" ", 1)[0]
    return first.count(".") >= 2 or (
        first.count(".") == 1 and len(first.split(".")) > 1 and first.split(".")[1].isdigit()
    )


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("Usage: format_toc_like_sample.py <docx-path>")

    path = Path(sys.argv[1])
    document = Document(path)

    inside_toc = False
    formatted = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == "TABLE OF CONTENTS":
            inside_toc = True
            continue
        if inside_toc and text.upper() == "ABSTRACT":
            break
        if not inside_toc or not text:
            continue

        title = text.split("\t", 1)[0].strip()
        subheading = is_subheading(title)

        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.22 if subheading else 0)
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(4)

        run = paragraph.add_run(title)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12 if subheading else 13)
        run.bold = not subheading
        formatted += 1

    document.save(path)
    print(f"formatted {formatted} TOC entries in {path}")


if __name__ == "__main__":
    main()
