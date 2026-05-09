from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


PAGE_MAP = {
    "ABSTRACT": 8,
    "1. PROBLEM DEFINITION": 9,
    "1.1 OVERVIEW": 9,
    "1.2 PROBLEM STATEMENT": 10,
    "2. INTRODUCTION": 11,
    "2.1 BACKGROUND OF THE PROJECT": 11,
    "2.2 OBJECTIVES OF THE PROJECT": 12,
    "2.3 SCOPE OF THE PROJECT": 12,
    "2.4 METHODOLOGY": 13,
    "3. LITERATURE SURVEY": 14,
    "3.1 REVIEW OF RESUME SCREENING METHODS": 14,
    "3.2 REVIEW OF APPLICANT TRACKING SYSTEMS": 15,
    "3.3 REVIEW OF NLP IN RESUME ANALYSIS": 15,
    "3.4 REVIEW OF JOB RECOMMENDATION TECHNIQUES": 16,
    "4. SYSTEM ANALYSIS": 17,
    "4.1 EXISTING SYSTEM": 17,
    "4.2 LIMITATIONS OF EXISTING SYSTEM": 18,
    "4.3 PROPOSED SYSTEM": 18,
    "4.4 FEASIBILITY STUDY": 19,
    "4.5 FUNCTIONAL REQUIREMENTS": 19,
    "4.6 NON-FUNCTIONAL REQUIREMENTS": 19,
    "5. SYSTEM DESIGN": 21,
    "5.1 SYSTEM ARCHITECTURE": 22,
    "5.2 DATA FLOW DIAGRAM DESCRIPTION": 22,
    "5.3 DATABASE DESIGN": 23,
    "5.4 USER INTERFACE DESIGN": 23,
    "5.5 SECURITY DESIGN": 23,
    "5.6 MODULE DESIGN": 24,
    "6. IMPLEMENTATION": 25,
    "6.1 TECHNOLOGY STACK AND PROJECT STRUCTURE": 26,
    "6.2 RESUME UPLOAD AND TEXT EXTRACTION": 28,
    "6.3 RESUME ANALYSIS, SCORING AND MATCHING": 29,
    "6.4 REPORT STORAGE, HISTORY AND PDF EXPORT": 31,
    "6.5 ADMIN, CONTACT AND VALIDATION MODULES": 33,
    "6.6 TESTING STRATEGY AND TEST CASE SUMMARY": 35,
    "7. RESULT": 37,
    "7.1 SAMPLE OUTPUT DESCRIPTION": 38,
    "7.2 PERFORMANCE OBSERVATION": 38,
    "7.3 ADVANTAGES OF THE SYSTEM": 39,
    "7.4 LIMITATIONS": 39,
    "7.5 FUTURE ENHANCEMENTS": 39,
    "8. CONCLUSION": 41,
    "8.1 SUMMARY OF WORK COMPLETED": 41,
    "8.2 LEARNING OUTCOMES": 41,
    "APPENDIX A: PROJECT FILE STRUCTURE": 43,
    "APPENDIX B: IMPORTANT ROUTES": 44,
    "APPENDIX C: IMPORTANT FUNCTIONS": 45,
    "APPENDIX D: REQUIREMENTS": 46,
    "REFERENCES": 47,
    "APPENDIX E: IMPORTANT CODE SNIPPETS": 48,
}


def normalize(title: str) -> str:
    return title.strip().upper()


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("Usage: update_toc_numbers_from_map.py <docx-path>")
    path = Path(sys.argv[1])
    document = Document(path)

    toc_start = next(
        (idx for idx, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "TABLE OF CONTENTS"),
        None,
    )
    if toc_start is None:
        raise SystemExit("TABLE OF CONTENTS not found")

    toc_end = next(
        (
            idx
            for idx, paragraph in enumerate(document.paragraphs[toc_start + 1 :], start=toc_start + 1)
            if paragraph.text.strip().upper() == "ABSTRACT" and paragraph.style.name.startswith("Heading")
        ),
        len(document.paragraphs),
    )

    updated = 0
    for paragraph in document.paragraphs[toc_start + 1 : toc_end]:
        text = paragraph.text.strip()
        if text:
            title = text.split("\t", 1)[0].strip()
            page = PAGE_MAP.get(normalize(title))
            if page is None:
                continue

            was_subheading = title[:3].count(".") >= 2 or title[:4].count(".") >= 2
            is_main = not was_subheading
            paragraph.clear()
            paragraph.paragraph_format.left_indent = Inches(0.28 if was_subheading else 0)
            paragraph.paragraph_format.line_spacing = 1.5
            run = paragraph.add_run(f"{title}\t{page}")
            run.bold = is_main
            run.font.name = "Times New Roman"
            run.font.size = Pt(17 if is_main else 10.5)
            updated += 1

    document.save(path)
    print(f"updated {updated} TOC entries in {path}")


if __name__ == "__main__":
    main()
