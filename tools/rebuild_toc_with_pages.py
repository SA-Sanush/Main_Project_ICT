from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


TOC_ENTRIES = [
    ("Abstract", 8, False),
    ("1. Problem Definition", 9, False),
    ("1.1 Overview", 9, True),
    ("1.2 Problem Statement", 10, True),
    ("2. Introduction", 11, False),
    ("2.1 Background Of The Project", 11, True),
    ("2.2 Objectives Of The Project", 12, True),
    ("2.3 Scope Of The Project", 12, True),
    ("2.4 Methodology", 13, True),
    ("3. Literature Survey", 14, False),
    ("3.1 Review Of Resume Screening Methods", 14, True),
    ("3.2 Review Of Applicant Tracking Systems", 15, True),
    ("3.3 Review Of NLP In Resume Analysis", 15, True),
    ("3.4 Review Of Job Recommendation Techniques", 16, True),
    ("4. System Analysis", 17, False),
    ("4.1 Existing System", 17, True),
    ("4.2 Limitations Of Existing System", 18, True),
    ("4.3 Proposed System", 18, True),
    ("4.4 Feasibility Study", 19, True),
    ("4.5 Functional Requirements", 19, True),
    ("4.6 Non-Functional Requirements", 19, True),
    ("5. System Design", 21, False),
    ("5.1 System Architecture", 22, True),
    ("5.2 Data Flow Diagram Description", 22, True),
    ("5.3 Database Design", 23, True),
    ("5.4 User Interface Design", 23, True),
    ("5.5 Security Design", 23, True),
    ("5.6 Module Design", 24, True),
    ("6. Implementation", 25, False),
    ("6.1 Technology Stack And Project Structure", 26, True),
    ("6.2 Resume Upload And Text Extraction", 28, True),
    ("6.3 Resume Analysis, Scoring And Matching", 29, True),
    ("6.4 Report Storage, History And Pdf Export", 31, True),
    ("6.5 Admin, Contact And Validation Modules", 33, True),
    ("6.6 Testing Strategy And Test Case Summary", 35, True),
    ("7. Result", 37, False),
    ("7.1 Sample Output Description", 38, True),
    ("7.2 Performance Observation", 38, True),
    ("7.3 Advantages Of The System", 39, True),
    ("7.4 Limitations", 39, True),
    ("7.5 Future Enhancements", 39, True),
    ("8. Conclusion", 41, False),
    ("8.1 Summary Of Work Completed", 41, True),
    ("8.2 Learning Outcomes", 41, True),
    ("Appendix A: Project File Structure", 43, False),
    ("Appendix B: Important Routes", 44, False),
    ("Appendix C: Important Functions", 45, False),
    ("Appendix D: Requirements", 46, False),
    ("References", 47, False),
    ("Appendix E: Important Code Snippets", 48, False),
]


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("Usage: rebuild_toc_with_pages.py <docx-path>")
    path = Path(sys.argv[1])
    document = Document(path)

    toc_start = next(
        (idx for idx, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "TABLE OF CONTENTS"),
        None,
    )
    if toc_start is None:
        raise SystemExit("TABLE OF CONTENTS not found")

    toc_end = next(
        (
            idx
            for idx, paragraph in enumerate(document.paragraphs[toc_start + 1 :], start=toc_start + 1)
            if paragraph.text.strip().upper() == "ABSTRACT" and paragraph.style.name.startswith("Heading")
        ),
        len(document.paragraphs),
    )

    toc_paragraphs = document.paragraphs[toc_start + 1 : toc_end]
    if len(toc_paragraphs) < len(TOC_ENTRIES):
        raise SystemExit("TOC has fewer paragraphs than expected")

    for paragraph, (title, page, subheading) in zip(toc_paragraphs, TOC_ENTRIES):
        paragraph.clear()
        paragraph.paragraph_format.left_indent = Inches(0.28 if subheading else 0)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(f"{title}\t{page}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12 if subheading else 13)
        run.bold = not subheading

    for paragraph in toc_paragraphs[len(TOC_ENTRIES) :]:
        paragraph.clear()

    document.save(path)
    print(f"rebuilt {len(TOC_ENTRIES)} TOC entries in {path}")


if __name__ == "__main__":
    main()
